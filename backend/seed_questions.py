"""
Seed script: parses all .docx question files in the project root.
Run from project root or backend/: python backend/seed_questions.py
"""
import sys, os, glob, re, unicodedata
sys.path.insert(0, os.path.join(os.path.dirname(__file__)))

import docx
from app.db.session import engine, SessionLocal
from app.db.base import Base
from app.models.curriculum import Topic
from app.models.assessment import Question

BASE_DIR = "/home/runner/workspace"

def nfc(s: str) -> str:
    return unicodedata.normalize("NFC", s) if s else s

def parse_filename(filename: str):
    name = os.path.basename(filename).replace(".docx", "")
    grade_match = re.search(r"(\d+)$", name)
    grade = int(grade_match.group(1)) if grade_match else None
    subject_raw = nfc(re.sub(r"\d+$", "", name).strip())
    subject_map = {nfc("Vật lý"): "Lý", "Hóa": "Hóa", "Lý": "Lý", "Sinh": "Sinh", "Toán": "Toán"}
    subject = subject_map.get(subject_raw, subject_raw)
    return subject, grade

def parse_docx(filepath):
    doc = docx.Document(filepath)
    topics = {}
    current_topic = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()

        if "heading" in style:
            current_topic = text.strip()
            if current_topic not in topics:
                topics[current_topic] = []
        elif current_topic:
            cleaned = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
            if cleaned and len(cleaned) > 8:
                topics[current_topic].append(cleaned)

    if not topics:
        current_topic = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            cleaned = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
            is_question = any(cleaned.startswith(w) for w in [
                "Viết","Tính","Cho","Xác","Tìm","Giải","Phân","So","Dự","Vẽ","Áp","Mắc",
                "Chứng","Tại","Nêu","Lấy","Phương","Hoàn","Liệt","Giải thích","Dùng","Một","Hai","Khi"
            ])
            if not is_question and len(text) < 100 and text == cleaned:
                current_topic = text
                if current_topic not in topics:
                    topics[current_topic] = []
            elif current_topic and is_question and len(cleaned) > 8:
                topics[current_topic].append(cleaned)

    return topics

def seed_file(db, filepath):
    subject, grade = parse_filename(filepath)
    if not grade:
        print(f"⚠️  Skipping {os.path.basename(filepath)}: no grade detected")
        return

    topics_data = parse_docx(filepath)
    total_q = sum(len(qs) for qs in topics_data.values())
    print(f"\n📚 {os.path.basename(filepath)} → {subject} lớp {grade} | {len(topics_data)} topics, {total_q} questions")

    for topic_name, questions in topics_data.items():
        if not topic_name or not questions:
            continue

        existing = db.query(Topic).filter(
            Topic.subject == subject,
            Topic.grade == grade,
            Topic.name == topic_name
        ).first()

        if not existing:
            topic = Topic(subject=subject, grade=grade, name=topic_name)
            db.add(topic)
            db.flush()
        else:
            topic = existing

        added = 0
        for q_text in questions:
            exists = db.query(Question).filter(
                Question.topic_id == topic.id,
                Question.content == q_text
            ).first()
            if not exists:
                db.add(Question(topic_id=topic.id, content=q_text, difficulty=0.5))
                added += 1

        if added:
            print(f"  ✓ {topic_name[:50]}: +{added} câu")

    db.commit()

def main():
    Base.metadata.create_all(bind=engine)
    db = SessionLocal()

    docx_files = glob.glob(os.path.join(BASE_DIR, "*.docx"))
    if not docx_files:
        print("❌ No .docx files found!")
        return

    print(f"Found {len(docx_files)} files")
    for filepath in sorted(docx_files):
        try:
            seed_file(db, filepath)
        except Exception as e:
            print(f"❌ Error in {filepath}: {e}")
            import traceback; traceback.print_exc()
            db.rollback()

    total_topics = db.query(Topic).count()
    total_questions = db.query(Question).count()
    print(f"\n✅ DONE! {total_topics} topics, {total_questions} questions seeded.")
    db.close()

if __name__ == "__main__":
    main()
